"""Geração do MOP em DOCX (Word)."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from ..models import MOP

PLACEHOLDER = "A preencher"


def _add_field(paragraph, field_code: str) -> None:
    """Insere um campo do Word (ex.: PAGE, NUMPAGES) em um parágrafo."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(begin)
    r.append(instr)
    r.append(end)


def _setup_header_footer(doc: Document, mop: MOP) -> None:
    """Configura cabeçalho (logo + texto) e rodapé (número de página).

    A primeira página (capa) não recebe cabeçalho/rodapé.
    """
    b = mop.branding
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Cabeçalho: logo à esquerda + texto.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if b.logo_path and os.path.isfile(b.logo_path):
        try:
            hp.add_run().add_picture(b.logo_path, height=Inches(0.35))
        except Exception:
            pass  # logo inválida: segue sem imagem
    if b.header_text:
        run = hp.add_run(("   " if hp.runs else "") + b.header_text)
        run.bold = True

    # Rodapé: número de página centralizado ("Página X de Y").
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Página ")
    _add_field(fp, "PAGE")
    fp.add_run(" de ")
    _add_field(fp, "NUMPAGES")


def _add_cover(doc: Document, mop: MOP) -> None:
    """Monta a capa: logo, título do projeto e capa da modificação."""
    b = mop.branding

    # Logo centralizada no topo.
    if b.logo_path and os.path.isfile(b.logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(b.logo_path, width=Inches(2.2))
        except Exception:
            pass

    # Espaçamento
    for _ in range(3):
        doc.add_paragraph()

    # Título do projeto.
    project_title = b.project_title or mop.title
    if project_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_title)
        run.bold = True
        run.font.size = Pt(20)

    # Capa da modificação (título formal do documento).
    cover_title = b.cover_title or mop.title
    if cover_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_title)
        run.font.size = Pt(16)

    if mop.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(mop.subtitle)
        run.italic = True
        run.font.size = Pt(12)

    doc.add_page_break()


def _set_update_fields_on_open(doc: Document) -> None:
    """Marca o documento para atualizar campos (incl. o índice) ao abrir.

    Assim o Word/LibreOffice recalcula o TOC automaticamente e o texto
    placeholder é substituído pelo índice real.
    """
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def _toc_entries(mop: MOP):
    """Entradas fixas do sumário (nível, título), conforme a estrutura do MOP."""
    entries = [
        (1, "1. Project Overview"),
        (2, "1.1 Objetivo"),
        (2, "1.2 Upgrade Plan"),
        (2, "1.3 Impacto"),
        (2, "1.4 Plano de Backup"),
        (2, "1.5 Validação Pós Implementação"),
        (2, "1.6 Plano de Volta (Rollback)"),
        (2, "1.7 Execução"),
        (2, "1.8 Acesso ao DevOps"),
    ]
    if mop.api_changes:
        entries.append((1, "2. Alterações de API (Endpoints e Payloads)"))
    if mop.git and (mop.git.commits or mop.git.changed_files):
        entries.append((1, "Anexo A — Detalhes do Git"))
    return entries


def _add_toc(doc: Document, mop: MOP) -> None:
    """Insere o sumário como um campo TOC com resultado estático embutido.

    O resultado estático (lista de seções) garante que Word, LibreOffice e até
    visualizadores simples mostrem o sumário mesmo sem atualizar o campo. Ao
    atualizar (F9/abrir no Word), o campo vira um índice dinâmico com páginas.
    """
    doc.add_heading("Sumário", level=1)
    entries = _toc_entries(mop)

    # Parágrafo de abertura do campo: begin + instrText + separate
    p_begin = doc.add_paragraph()
    run = p_begin.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)

    # Resultado estático: uma linha por seção (indentada por nível).
    for level, title in entries:
        p = doc.add_paragraph(title)
        p.paragraph_format.left_indent = Inches(0.3 * (level - 1))

    # Parágrafo de fechamento do campo: end
    p_end = doc.add_paragraph()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p_end.add_run()._r.append(end)


def _add_paragraph(doc: Document, text: str) -> None:
    """Adiciona um parágrafo; se vazio, insere um placeholder em itálico."""
    if text:
        doc.add_paragraph(text)
    else:
        p = doc.add_paragraph()
        run = p.add_run(PLACEHOLDER)
        run.italic = True


def _add_bold_label(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    p.add_run(label).bold = True


def _add_code_block(doc: Document, text: str) -> None:
    """Adiciona um bloco de código monoespaçado, preservando quebras de linha."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def _add_api_section(doc: Document, mop: MOP) -> None:
    """Renderiza a seção de alterações de API (endpoints, payloads, headers)."""
    if not mop.api_changes:
        return
    doc.add_heading("2. Alterações de API (Endpoints e Payloads)", level=1)
    for ep in mop.api_changes:
        doc.add_heading(ep.title(), level=2)
        if ep.description:
            doc.add_paragraph(ep.description)

        if ep.request_headers:
            _add_bold_label(doc, "Headers (request):")
            for h in ep.request_headers:
                doc.add_paragraph(h, style="List Bullet")
        if ep.request_body:
            _add_bold_label(doc, "Corpo da requisição (request):")
            _add_code_block(doc, ep.request_body)

        if ep.response_status:
            _add_bold_label(doc, f"Resposta (response) — {ep.response_status}:")
        elif ep.response_body or ep.response_headers:
            _add_bold_label(doc, "Resposta (response):")
        if ep.response_headers:
            for h in ep.response_headers:
                doc.add_paragraph(h, style="List Bullet")
        if ep.response_body:
            _add_code_block(doc, ep.response_body)

        if ep.notes:
            p = doc.add_paragraph()
            p.add_run("Notas: ").bold = True
            p.add_run(ep.notes)


def generate_docx(mop: MOP, output_path: str) -> None:
    """Renderiza o MOP em um arquivo .docx no caminho informado."""
    doc = Document()

    # Cabeçalho (logo + texto) e rodapé (número de página); capa sem eles.
    _setup_header_footer(doc, mop)

    # Capa da modificação (logo, título do projeto, título do documento).
    _add_cover(doc, mop)

    # Sumário logo após a capa.
    _add_toc(doc, mop)
    doc.add_page_break()

    # 1. Project Overview
    doc.add_heading("1. Project Overview", level=1)

    # 1.1 Objetivo
    doc.add_heading("1.1 Objetivo", level=2)
    _add_paragraph(doc, mop.objetivo)

    if mop.mudancas:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Id"
        hdr[1].text = "Ação"
        hdr[2].text = "Descrição"
        hdr[3].text = "Notas"
        for idx, ch in enumerate(mop.mudancas, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = ch.acao
            row[2].text = ch.descricao
            row[3].text = ch.notas
    else:
        _add_paragraph(doc, "")

    # 1.2 Upgrade Plan
    doc.add_heading("1.2 Upgrade Plan", level=2)
    p = doc.add_paragraph()
    p.add_run("Upgrade Schedule: ").bold = True
    p.add_run(mop.upgrade_schedule or PLACEHOLDER)

    # 1.3 Impacto
    doc.add_heading("1.3 Impacto", level=2)
    if mop.impacto:
        for item in mop.impacto:
            doc.add_paragraph(item, style="List Bullet")
    else:
        _add_paragraph(doc, "")

    # 1.4 Plano de Backup
    doc.add_heading("1.4 Plano de Backup", level=2)
    _add_paragraph(doc, mop.plano_backup)
    for pr in mop.pull_requests:
        doc.add_paragraph(pr, style="List Bullet")

    # 1.5 Validação Pós Implementação
    doc.add_heading("1.5 Validação Pós Implementação", level=2)
    _add_paragraph(doc, mop.validacao)

    # 1.6 Plano de Volta (Rollback)
    doc.add_heading("1.6 Plano de Volta (Rollback)", level=2)
    _add_paragraph(doc, mop.rollback)
    for pr in mop.pull_requests:
        doc.add_paragraph(pr, style="List Bullet")

    # 1.7 Execução
    doc.add_heading("1.7 Execução", level=2)
    rt = mop.responsavel_tecnico
    p = doc.add_paragraph()
    p.add_run("Responsável Técnico (Quem vai executar)").bold = True
    if rt.empresa:
        doc.add_paragraph(rt.empresa)
    for person in rt.pessoas:
        doc.add_paragraph(person.render(), style="List Bullet")
    if not rt.empresa and not rt.pessoas:
        doc.add_paragraph(PLACEHOLDER, style="List Bullet")

    p = doc.add_paragraph()
    p.add_run(
        "Equipe Day-After (Quem irá acompanhar o ambiente no dia posterior)"
    ).bold = True
    da = mop.day_after
    if da.empresa:
        doc.add_paragraph(da.empresa)
    if da.papel:
        doc.add_paragraph(da.papel, style="List Bullet")
    if not da.empresa and not da.papel:
        doc.add_paragraph(PLACEHOLDER, style="List Bullet")

    # 1.8 Acesso ao DevOps
    doc.add_heading("1.8 Acesso ao DevOps", level=2)
    if mop.acesso_devops:
        _add_paragraph(doc, f"Necessário o acesso ao DevOps: {mop.acesso_devops}")
    else:
        _add_paragraph(doc, "")

    # 2. Alterações de API (somente se houver mudança de endpoints/payloads)
    _add_api_section(doc, mop)

    # Anexo: detalhes do Git
    if mop.git and (mop.git.commits or mop.git.changed_files):
        g = mop.git
        doc.add_heading("Anexo A — Detalhes do Git", level=1)
        doc.add_paragraph(f"Repositório: {g.repo_url}")
        doc.add_paragraph(f"Branch: {g.branch}")
        if g.base:
            doc.add_paragraph(f"Base: {g.base}")

        if g.commits:
            doc.add_heading(f"Commits ({len(g.commits)})", level=2)
            for c in g.commits:
                doc.add_paragraph(
                    f"{c.short_sha} {c.subject} — {c.author} ({c.date})",
                    style="List Bullet",
                )
        if g.changed_files:
            doc.add_heading(f"Arquivos alterados ({len(g.changed_files)})", level=2)
            for filepath in g.changed_files:
                doc.add_paragraph(filepath, style="List Bullet")

    _set_update_fields_on_open(doc)
    doc.save(output_path)
