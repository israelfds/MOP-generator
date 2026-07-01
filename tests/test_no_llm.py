"""Testes do fluxo SEM LLM (offline, usando um repositório Git real)."""

from __future__ import annotations

import os

from click.testing import CliRunner

from mop_generator.cli import cli
from mop_generator.collector import build_mop, pending_fields
from mop_generator.config import mop_from_config
from mop_generator.generators import generate_docx, generate_markdown
from mop_generator.git_utils import collect_git_info, prepare_repo, ssh_to_web_url


def test_collect_git_info(git_repo):
    repo, branch, base = git_repo
    with prepare_repo(repo, branch) as path:
        info = collect_git_info(path, repo, branch, base)

    assert info.branch == branch
    assert info.base == base
    # 2 commits na feature em relação a main
    assert len(info.commits) == 2
    subjects = [c.subject for c in info.commits]
    assert "feat: adiciona pipeline de correcao" in subjects
    assert set(info.changed_files) == {"pipeline.yaml", "tasks.py"}
    assert info.diff_text  # capturou algum diff


def test_build_mop_derives_title_and_changes_from_commits(git_repo):
    repo, branch, base = git_repo
    with prepare_repo(repo, branch) as path:
        info = collect_git_info(path, repo, branch, base)

    # sem config, não interativo: sugere mudanças a partir dos commits
    mop = build_mop(mop_from_config({}), info, interactive=False)

    assert mop.title == "MOP Upgrade repo"  # nome do diretório do repo
    assert len(mop.mudancas) == 2  # uma por commit
    # campos humanos ficam pendentes
    pend = pending_fields(mop)
    assert "Objetivo" in pend
    assert "Responsável Técnico" in pend


def test_markdown_generation_with_config(git_repo):
    repo, branch, base = git_repo
    cfg = {
        "title": "MOP Teste",
        "objetivo": "Objetivo de teste",
        "upgrade_schedule": "Janela 01/01",
        "impacto": ["Impacto A"],
        "branding": {"cover_title": "MOP Teste Capa", "project_title": "Projeto X"},
        "responsavel_tecnico": {
            "empresa": "Acme",
            "pessoas": [{"nome": "Fulano", "email": "f@x.com"}],
        },
        "day_after": {"empresa": "Ops Team", "papel": "Plantonista"},
    }
    with prepare_repo(repo, branch) as path:
        info = collect_git_info(path, repo, branch, base)
    mop = build_mop(mop_from_config(cfg), info, interactive=False)
    md = generate_markdown(mop)

    assert "# MOP Teste Capa" in md
    assert "Projeto X" in md
    assert "Objetivo de teste" in md
    assert "Fulano <f@x.com>" in md
    assert "Janela 01/01" in md
    # sem PRs, deve aparecer placeholder em pendências (backup/rollback vazios)
    assert "A preencher" in md


def test_placeholders_for_empty_fields(git_repo):
    repo, branch, base = git_repo
    with prepare_repo(repo, branch) as path:
        info = collect_git_info(path, repo, branch, base)
    mop = build_mop(mop_from_config({}), info, interactive=False)
    md = generate_markdown(mop)
    # objetivo vazio deve render placeholder
    assert "A preencher" in md


def test_docx_generation(tmp_path, git_repo):
    from docx import Document

    repo, branch, base = git_repo
    with prepare_repo(repo, branch) as path:
        info = collect_git_info(path, repo, branch, base)
    mop = build_mop(mop_from_config({"objetivo": "obj"}), info, interactive=False)

    out = tmp_path / "out.docx"
    generate_docx(mop, str(out))
    assert out.exists()

    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "1.1 Objetivo" in texts
    assert "1.8 Acesso ao DevOps" in texts
    assert len(doc.tables) == 1  # tabela de mudanças


def test_ssh_to_web_url_azure():
    ssh = "git@ssh.dev.azure.com:v3/Org/Projeto/repo"
    assert ssh_to_web_url(ssh) == "https://dev.azure.com/Org/Projeto/_git/repo"


def test_local_repo_shows_remote_name_not_path(git_repo, tmp_path):
    """Repo local com remote deve exibir a URL do remote, não o path do disco."""
    repo, branch, base = git_repo
    # adiciona um remote fake estilo Azure
    import subprocess

    subprocess.run(
        [
            "git",
            "remote",
            "add",
            "origin",
            "git@ssh.dev.azure.com:v3/Org/Proj/meurepo",
        ],
        cwd=repo,
        check=True,
        capture_output=True,
    )
    with prepare_repo(repo, branch) as path:
        info = collect_git_info(path, repo, branch, base)
    assert info.repo_url == "https://dev.azure.com/Org/Proj/_git/meurepo"
    assert not info.repo_url.startswith("/")  # não é path local


def test_cli_no_llm_end_to_end(tmp_path, git_repo):
    repo, branch, base = git_repo
    out = tmp_path / "mop.md"
    runner = CliRunner()
    result = runner.invoke(
        cli,
        [
            "generate",
            "--repo",
            repo,
            "--branch",
            branch,
            "--base",
            base,
            "--format",
            "md",
            "--no-llm",
            "--non-interactive",
            "-o",
            str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    assert out.exists()
    content = out.read_text(encoding="utf-8")
    assert "1. Project Overview" in content
    # sem LLM e sem config, deve avisar sobre pendências
    assert "pendentes" in result.output.lower()


def test_project_context_detection(tmp_path):
    """Detecta linguagem, stack e nome/descrição de forma agnóstica."""
    import subprocess

    from mop_generator.project_context import detect_project_context

    repo = tmp_path / "svc"
    repo.mkdir()
    (repo / "requirements.txt").write_text("fastapi\ncelery\n", encoding="utf-8")
    (repo / "Dockerfile").write_text("FROM python:3.11\n", encoding="utf-8")
    (repo / "main.py").write_text("print('hi')\n", encoding="utf-8")
    (repo / "README.md").write_text(
        "# Meu Servico\n\nServico de exemplo para registro de usuarios.\n",
        encoding="utf-8",
    )

    ctx = detect_project_context(str(repo), changed_files=["main.py"], repo_name="svc")
    assert "Python" in ctx.languages
    assert "Docker" in ctx.stack
    assert "FastAPI" in ctx.stack
    assert "Celery" in ctx.stack
    assert ctx.name == "Meu Servico"
    assert "registro de usuarios" in ctx.description


def test_docx_has_toc(tmp_path, git_repo):
    from docx import Document

    repo, branch, base = git_repo
    with prepare_repo(repo, branch) as path:
        info = collect_git_info(path, repo, branch, base)
    mop = build_mop(mop_from_config({"objetivo": "obj"}), info, interactive=False)

    out = tmp_path / "out.docx"
    generate_docx(mop, str(out))
    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Sumário" in texts  # heading do índice
    # a seção "Contexto do projeto" NÃO deve aparecer no documento
    assert "Contexto do projeto" not in texts
    # o campo TOC foi inserido no XML
    assert "TOC" in doc.element.xml


def test_docx_cover_header_footer(tmp_path, git_repo):
    """Capa com título, cabeçalho com texto e rodapé com número de página."""
    from docx import Document

    from mop_generator.models import Branding

    repo, branch, base = git_repo
    with prepare_repo(repo, branch) as path:
        info = collect_git_info(path, repo, branch, base)
    mop = build_mop(mop_from_config({"objetivo": "obj"}), info, interactive=False)
    mop.branding = Branding(
        project_title="Sistema ABC",
        cover_title="MOP para Atualização em Produção – ABC",
        header_text="Minha Empresa",
    )

    out = tmp_path / "cover.docx"
    generate_docx(mop, str(out))
    doc = Document(str(out))

    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Sistema ABC" in texts
    assert "MOP para Atualização em Produção – ABC" in texts

    section = doc.sections[0]
    assert section.different_first_page_header_footer is True
    header_text = "\n".join(p.text for p in section.header.paragraphs)
    assert "Minha Empresa" in header_text
    # rodapé deve conter campo de página
    assert "PAGE" in section.footer.part.element.xml


def test_list_branches(git_repo):
    from mop_generator.git_utils import current_branch, list_branches

    repo, branch, base = git_repo
    with prepare_repo(repo, branch) as path:
        branches = list_branches(path)
        cur = current_branch(path)
    assert "main" in branches
    assert "feature/registro" in branches
    assert cur == "feature/registro"


def test_core_generate_mop_no_llm(tmp_path, git_repo):
    from mop_generator.core import generate_mop

    repo, branch, base = git_repo
    out = tmp_path / "core.md"
    result = generate_mop(
        repo=repo,
        branch=branch,
        base=base,
        fmt="md",
        output=str(out),
        interactive=False,
        use_llm=False,
    )
    assert result.output_path == str(out)
    assert out.exists()
    assert "Objetivo" in result.pending  # sem config/LLM
    assert result.used_llm is False
